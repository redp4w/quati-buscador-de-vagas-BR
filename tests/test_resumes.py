from datetime import UTC, datetime
from io import BytesIO

import pytest
from docx import Document
from pypdf import PdfReader

from quati.ai import analyze_locally
from quati.applications import prepare_application
from quati.domain import JobRecord
from quati.profile import CandidateProfile
from quati.resumes import (
    ResumeDocument,
    ResumeVault,
    build_external_resume_prompt,
    create_cover_letter,
    export_docx,
    export_html,
    export_pdf,
    extract_resume,
    profile_fields_from_resume,
    resume_from_profile,
    resume_section_titles,
    tailor_document,
)


def _docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Pessoa Desenvolvedora com experiência em Python e SQL.")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _profile_docx_bytes() -> bytes:
    document = Document()
    for line in (
        "Ana Souza",
        "Analista de Segurança da Informação",
        "ana@example.com | (11) 99999-1234",
        "Resumo profissional",
        "Experiência com gestão de vulnerabilidades.",
        "Competências",
        "Python, SIEM, ISO 27001",
        "Experiência profissional",
        "Analista de SOC — Empresa Exemplo",
        "Formação acadêmica",
        "Tecnologia em Segurança da Informação",
        "https://www.linkedin.com/in/ana-souza",
    ):
        document.add_paragraph(line)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _job() -> JobRecord:
    now = datetime.now(UTC)
    return JobRecord(
        id=1,
        source="gupy",
        external_id="1",
        title="Pessoa Desenvolvedora",
        company="Acme",
        location="Remoto",
        url="https://acme.gupy.io/jobs/1",
        description="Requisitos: Python e SQL.",
        published_at="",
        status="active",
        first_seen_at=now,
        last_seen_at=now,
    )


def test_extracts_docx_in_memory_and_exports_local_pdf() -> None:
    resume = extract_resume("curriculo.docx", _docx_bytes())
    document = tailor_document(resume, _job(), analyze_locally(_job(), resume.text))
    pdf = export_pdf("Currículo", document)

    assert "Python" in resume.text
    assert pdf.startswith(b"%PDF")


def test_resume_prefills_profile_fields_without_saving() -> None:
    resume = extract_resume("curriculo.docx", _profile_docx_bytes())
    fields = profile_fields_from_resume(resume)

    assert "\n" in resume.text
    assert fields["name"] == "Ana Souza"
    assert fields["headline"] == "Analista de Segurança da Informação"
    assert fields["email"] == "ana@example.com"
    assert "99999-1234" in fields["phone"]
    assert fields["skills"] == "Python, SIEM, ISO 27001"
    assert "Analista de SOC" in fields["experience"]
    assert "linkedin.com/in/ana-souza" in fields["links"]


@pytest.mark.parametrize("filename", ["curriculo.txt", ""])
def test_rejects_unsupported_or_missing_files(filename: str) -> None:
    with pytest.raises(ValueError):
        extract_resume(filename, b"conteudo")


def test_resume_library_encrypts_imports_and_supports_deletion(tmp_path) -> None:
    path = tmp_path / "resumes.enc"
    vault = ResumeVault(path)
    content = _docx_bytes()
    stored = vault.add(
        "senha-local-segura", label="Backend", filename="curriculo.docx", content=content
    )

    loaded = vault.load("senha-local-segura")
    assert loaded[0].id == stored.id
    assert loaded[0].content == content
    assert b"Backend" not in path.read_bytes()
    assert b"Python" not in path.read_bytes()

    with pytest.raises(ValueError, match="já está"):
        vault.add(
            "senha-local-segura", label="Duplicado", filename="curriculo.docx", content=content
        )
    with pytest.raises(ValueError):
        vault.load("senha-incorreta")

    vault.delete("senha-local-segura", stored.id)
    assert vault.load("senha-local-segura") == []


def test_profile_can_prepare_pdf_and_docx_application_bundle() -> None:
    profile = CandidateProfile(
        "Ana", "ana@example.com", "", "", "Python, SQL", "ADS", "Desenvolvimento Python"
    )
    bundle = prepare_application(_job(), resume_from_profile(profile), tailor=True)

    assert bundle.analysis.compatibility_score > 0
    assert bundle.pdf.startswith(b"%PDF")
    assert bundle.docx.startswith(b"PK")


def test_cover_letter_uses_real_resume_and_job_context() -> None:
    resume = ResumeDocument(
        "curriculo.docx",
        """Ana Souza
Analista de Segurança

Resumo profissional
Experiência com monitoramento de eventos e resposta a incidentes.

Competências
SIEM, Python, ISO 27001
""",
    )
    job = _job()
    letter = create_cover_letter(resume, job)

    assert "Ana Souza" in letter
    assert job.title in letter
    assert job.company in letter
    assert "Experiência com monitoramento" in letter
    assert "experiências relevantes" not in letter


def test_exported_resume_has_professional_local_layout_and_no_author_metadata() -> None:
    body = """Ana Souza
Analista de Segurança da Informação
ana@example.com | (11) 99999-1234 | Itu, SP

Resumo profissional
Experiência com gestão de vulnerabilidades.

Competências
- SIEM
- ISO 27001

Experiência
Analista de SOC - Empresa Exemplo
"""

    docx = Document(BytesIO(export_docx("Currículo", body)))
    pdf_bytes = export_pdf("Currículo", body)
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_bytes)).pages)

    assert docx.paragraphs[0].text == "Ana Souza"
    assert docx.paragraphs[0].alignment is not None
    assert "Resume Section" in docx.styles
    assert docx.core_properties.author == ""
    assert "resumo profissional" in pdf_text.lower()
    assert "Ana Souza" in pdf_text
    assert not (PdfReader(BytesIO(pdf_bytes)).metadata.author or "")


def test_external_prompt_removes_identity_and_contacts() -> None:
    resume = extract_resume("curriculo.docx", _profile_docx_bytes())
    prompt = build_external_resume_prompt(resume, _job())

    assert "Ana Souza" not in prompt
    assert "ana@example.com" not in prompt
    assert "99999-1234" not in prompt
    assert "linkedin.com/in/ana-souza" not in prompt
    assert "gestão de vulnerabilidades" in prompt


def test_html_export_escapes_content_and_pdf_has_no_active_code() -> None:
    body = (
        "Pessoa Teste\n\nResumo profissional\n"
        "<script>alert('x')</script> #read(\"segredo-local.txt\")"
    )
    html = export_html("Currículo", body)
    pdf = export_pdf("Currículo", body)
    text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf)).pages)

    assert "<script>" not in html
    assert "&lt;script&gt;" in html
    assert "javascript:" not in html.lower()
    assert "segredo-local.txt" in text
    assert "Pessoa Teste" in text
    assert pdf.startswith(b"%PDF")


def test_html_export_supports_section_order_and_visual_options() -> None:
    body = """Ana Souza
Analista
ana@example.com

Resumo profissional
Resumo verdadeiro.

Competências
- Python
- SQL
"""

    assert resume_section_titles(body) == ("Resumo profissional", "Competências")
    html = export_html(
        "Currículo",
        body,
        template="classic",
        accent="graphite",
        density="compact",
        section_order=("Competências",),
    )

    assert '<body class="classic">' in html
    assert "#333333" in html
    assert "Competências" in html
    assert "Resumo verdadeiro" not in html


@pytest.mark.parametrize("template", ["contemporary", "executive"])
@pytest.mark.parametrize("accent", ["forest", "wine"])
def test_html_export_supports_additional_professional_templates(
    template: str, accent: str
) -> None:
    html = export_html(
        "Currículo",
        "Ana Souza\n\nResumo profissional\nExperiência verdadeira.",
        template=template,
        accent=accent,
    )

    assert f'<body class="{template}">' in html


def test_profile_parser_preserves_projects_additional_and_keywords() -> None:
    resume = ResumeDocument(
        "curriculo.pdf",
        """Ana Souza
Analista de Segurança
Itu - SP | Brasil • ana@example.com
Resumo profissional
Resumo.
Projetos acadêmicos e laboratoriais
Projeto de observabilidade
Experiência profissional
Analista - Empresa
Experiência anterior
- Atendimento
Diferenciais
- Perfil analítico
Palavras-chave ATS
SOC, SIEM, Python
""",
    )

    fields = profile_fields_from_resume(resume)

    assert fields["address"] == "Itu - SP | Brasil"
    assert fields["projects"] == "Projeto de observabilidade"
    assert "Experiência anterior" in fields["experience"]
    assert fields["additional"] == "- Perfil analítico"
    assert fields["keywords"] == "SOC, SIEM, Python"
